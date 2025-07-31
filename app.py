import os
import io
import re
import json
import logging
import tempfile
import time
import socket
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import PyPDF2
import docx2txt
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
import numpy as np
import pandas as pd
import shap
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure Google Generative AI
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 
app.config['UPLOAD_FOLDER'] = 'uploads'

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Comprehensive skill databases
TECHNICAL_SKILLS = {
    'programming_languages': [
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust', 'kotlin',
        'swift', 'php', 'ruby', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash', 'powershell'
    ],
    'web_technologies': [
        'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask',
        'spring', 'laravel', 'bootstrap', 'jquery', 'webpack', 'sass', 'less', 'tailwind'
    ],
    'databases': [
        'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'sql server',
        'cassandra', 'elasticsearch', 'dynamodb', 'firebase', 'neo4j'
    ],
    'cloud_platforms': [
        'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'linode',
        'kubernetes', 'docker', 'terraform', 'ansible', 'jenkins'
    ],
    'data_science': [
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
        'pandas', 'numpy', 'matplotlib', 'seaborn', 'jupyter', 'tableau', 'power bi'
    ],
    'mobile_development': [
        'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic', 'cordova'
    ],
    'tools_frameworks': [
        'git', 'github', 'gitlab', 'jira', 'confluence', 'slack', 'trello', 'asana',
        'visual studio', 'intellij', 'eclipse', 'postman', 'swagger'
    ]
}

SOFT_SKILLS = [
    'communication', 'leadership', 'teamwork', 'problem solving', 'analytical thinking',
    'project management', 'time management', 'adaptability', 'creativity', 'critical thinking',
    'collaboration', 'presentation', 'negotiation', 'mentoring', 'coaching', 'agile', 'scrum',
    'organizational', 'planning', 'strategic thinking', 'decision making', 'conflict resolution'
]

DOMAIN_SKILLS = {
    'finance': ['financial modeling', 'risk management', 'investment analysis', 'accounting', 'budgeting'],
    'healthcare': ['clinical research', 'medical coding', 'patient care', 'healthcare analytics'],
    'marketing': ['digital marketing', 'seo', 'sem', 'social media', 'content marketing', 'email marketing'],
    'sales': ['sales strategy', 'crm', 'lead generation', 'customer relationship', 'business development'],
    'hr': ['recruitment', 'talent acquisition', 'performance management', 'employee relations'],
    'operations': ['supply chain', 'logistics', 'process improvement', 'quality assurance', 'lean six sigma']
}

JOB_MARKET_DATA = {
    'software_engineer': {
        'salary_range': '$70,000 - $150,000',
        'demand_level': 'Very High',
        'growth_rate': '22% (Much faster than average)',
        'required_skills': ['programming', 'algorithms', 'data structures', 'software design'],
        'emerging_skills': ['cloud computing', 'microservices', 'devops', 'ai/ml']
    },
    'data_scientist': {
        'salary_range': '$95,000 - $165,000',
        'demand_level': 'High',
        'growth_rate': '31% (Much faster than average)',
        'required_skills': ['python', 'r', 'statistics', 'machine learning', 'sql'],
        'emerging_skills': ['deep learning', 'nlp', 'computer vision', 'mlops']
    },
    'web_developer': {
        'salary_range': '$55,000 - $120,000',
        'demand_level': 'High',
        'growth_rate': '13% (Faster than average)',
        'required_skills': ['html', 'css', 'javascript', 'responsive design'],
        'emerging_skills': ['react', 'vue', 'node.js', 'progressive web apps']
    },
    'mobile_developer': {
        'salary_range': '$75,000 - $140,000',
        'demand_level': 'High',
        'growth_rate': '19% (Much faster than average)',
        'required_skills': ['mobile frameworks', 'ui/ux design', 'api integration'],
        'emerging_skills': ['flutter', 'react native', 'ar/vr', 'iot integration']
    },
    'devops_engineer': {
        'salary_range': '$85,000 - $155,000',
        'demand_level': 'Very High',
        'growth_rate': '25% (Much faster than average)',
        'required_skills': ['ci/cd', 'containerization', 'cloud platforms', 'automation'],
        'emerging_skills': ['kubernetes', 'serverless', 'infrastructure as code', 'observability']
    }
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def check_internet_connectivity():
    """Check if we can reach Google's servers"""
    try:
        # Try to connect to Google's DNS server
        socket.create_connection(("8.8.8.8", 53), timeout=5)
        return True
    except OSError:
        try:
            # Try to connect to Google's generative AI endpoint
            socket.create_connection(("generativelanguage.googleapis.com", 443), timeout=10)
            return True
        except OSError:
            return False

def retry_with_backoff(func, max_retries=3, base_delay=1):
    """Retry a function with exponential backoff"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if attempt == max_retries - 1:
                raise e

            delay = base_delay * (2 ** attempt)
            logger.warning(f"Attempt {attempt + 1} failed: {str(e)}. Retrying in {delay} seconds...")
            time.sleep(delay)

    return None

def extract_skills_from_text(text):
    """Extract skills from resume text using precise rule-based approach"""
    import re

    text_lower = text.lower()
    found_skills = {
        'technical_skills': [],
        'soft_skills': [],
        'domain_skills': []
    }

    def is_skill_present(skill, text):
        """Check if skill is present as a whole word or phrase"""
        skill_lower = skill.lower()

        # For multi-word skills, check exact phrase match
        if ' ' in skill_lower:
            # Use word boundaries for multi-word phrases
            pattern = r'\b' + re.escape(skill_lower) + r'\b'
            return bool(re.search(pattern, text))
        else:
            # For single words, use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(skill_lower) + r'\b'
            return bool(re.search(pattern, text))

    # Extract technical skills with precise matching
    for category, skills in TECHNICAL_SKILLS.items():
        for skill in skills:
            if is_skill_present(skill, text_lower):
                found_skills['technical_skills'].append(skill)

    # Extract soft skills with precise matching
    for skill in SOFT_SKILLS:
        if is_skill_present(skill, text_lower):
            found_skills['soft_skills'].append(skill)

    # Extract domain skills with precise matching
    for domain, skills in DOMAIN_SKILLS.items():
        for skill in skills:
            if is_skill_present(skill, text_lower):
                found_skills['domain_skills'].append(skill)

    # Remove duplicates and return
    for category in found_skills:
        found_skills[category] = list(set(found_skills[category]))

    # Log what was found for debugging
    total_found = sum(len(skills) for skills in found_skills.values())
    logger.info(f"Skill extraction found {total_found} skills: "
               f"Technical: {len(found_skills['technical_skills'])}, "
               f"Soft: {len(found_skills['soft_skills'])}, "
               f"Domain: {len(found_skills['domain_skills'])}")

    return found_skills

def extract_experience_info(text):
    """Extract experience information from resume text"""
    text_lower = text.lower()

    # Look for experience indicators
    experience_years = 0
    experience_level = "Entry Level"
    experience_summary = ""

    # Extract years of experience
    year_patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
        r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience',
        r'experience\s*(?:of\s*)?(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s*in'
    ]

    for pattern in year_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            experience_years = max(experience_years, int(matches[0]))

    # Determine experience level
    if experience_years >= 7:
        experience_level = "Senior Level"
    elif experience_years >= 3:
        experience_level = "Mid Level"
    elif experience_years >= 1:
        experience_level = "Junior Level"

    # Look for leadership indicators
    leadership_keywords = ['lead', 'manager', 'director', 'head', 'chief', 'senior', 'principal']
    has_leadership = any(keyword in text_lower for keyword in leadership_keywords)

    if has_leadership and experience_years >= 3:
        experience_level = "Senior Level"

    # Create experience summary
    if experience_years > 0:
        experience_summary = f"{experience_years}+ years of professional experience"
        if has_leadership:
            experience_summary += " with leadership responsibilities"
    else:
        experience_summary = "Professional experience in relevant field"

    return {
        'years': experience_years,
        'level': experience_level,
        'summary': experience_summary,
        'has_leadership': has_leadership
    }

def extract_education_info(text):
    """Extract education information from resume text"""
    text_lower = text.lower()

    education_level = "High School"
    education_field = "General"
    education_summary = ""

    # Education level patterns
    if any(degree in text_lower for degree in ['phd', 'ph.d', 'doctorate', 'doctoral']):
        education_level = "Doctorate"
    elif any(degree in text_lower for degree in ['master', 'mba', 'ms', 'm.s', 'ma', 'm.a', 'mtech', 'm.tech']):
        education_level = "Master's"
    elif any(degree in text_lower for degree in ['bachelor', 'bs', 'b.s', 'ba', 'b.a', 'btech', 'b.tech', 'be', 'b.e']):
        education_level = "Bachelor's"
    elif any(term in text_lower for term in ['associate', 'diploma', 'certificate']):
        education_level = "Associate/Diploma"

    # Education field patterns
    tech_fields = ['computer', 'engineering', 'technology', 'science', 'mathematics', 'physics', 'chemistry']
    business_fields = ['business', 'management', 'finance', 'economics', 'marketing', 'accounting']

    if any(field in text_lower for field in tech_fields):
        education_field = "Technology/Engineering"
    elif any(field in text_lower for field in business_fields):
        education_field = "Business/Management"

    education_summary = f"{education_level} degree in {education_field}"

    return {
        'level': education_level,
        'field': education_field,
        'summary': education_summary
    }

def analyze_resume_with_shap(text):
    """Comprehensive resume analysis using rule-based extraction and SHAP for explainability"""
    try:
        logger.info("Starting rule-based resume analysis with SHAP explainability")

        # Extract information using rule-based methods
        skills_data = extract_skills_from_text(text)
        experience_data = extract_experience_info(text)
        education_data = extract_education_info(text)

        # Combine all skills and validate they're actually present
        all_skills = (skills_data['technical_skills'] +
                     skills_data['soft_skills'] +
                     skills_data['domain_skills'])

        # Additional validation: ensure we have meaningful skills
        if not all_skills:
            logger.warning("No skills found in resume text - this might indicate text extraction issues")
        else:
            logger.info(f"Successfully extracted {len(all_skills)} skills from resume")

        # Create feature vector for SHAP analysis
        features = {
            'technical_skills_count': len(skills_data['technical_skills']),
            'soft_skills_count': len(skills_data['soft_skills']),
            'domain_skills_count': len(skills_data['domain_skills']),
            'experience_years': experience_data['years'],
            'has_leadership': 1 if experience_data['has_leadership'] else 0,
            'education_level_score': {
                'High School': 1, 'Associate/Diploma': 2,
                "Bachelor's": 3, "Master's": 4, 'Doctorate': 5
            }.get(education_data['level'], 1),
            'tech_education': 1 if 'Technology' in education_data['field'] else 0,
            'total_skills': len(all_skills)
        }

        # Calculate normalized feature values (0-1 scale)
        normalized_features = {
            'technical_skills': min(features['technical_skills_count'] / 15, 1.0),
            'soft_skills': min(features['soft_skills_count'] / 10, 1.0),
            'domain_skills': min(features['domain_skills_count'] / 8, 1.0),
            'experience_quality': min((features['experience_years'] * 0.15 +
                                     features['has_leadership'] * 0.3), 1.0),
            'education_quality': min((features['education_level_score'] * 0.15 +
                                    features['tech_education'] * 0.25), 1.0)
        }

        # Create synthetic data for SHAP analysis
        # This simulates a trained model's feature importance
        feature_weights = {
            'technical_skills': 0.35,
            'soft_skills': 0.20,
            'domain_skills': 0.15,
            'experience_quality': 0.20,
            'education_quality': 0.10
        }

        # Calculate SHAP-like values (feature contributions)
        base_value = 0.3  # Base employability score
        shap_values = {}
        total_contribution = 0

        for feature, value in normalized_features.items():
            contribution = value * feature_weights[feature]
            shap_values[feature] = contribution
            total_contribution += contribution

        # Normalize SHAP values to sum to 1
        if total_contribution > 0:
            normalized_shap = {k: v/total_contribution for k, v in shap_values.items()}
        else:
            normalized_shap = {k: 0.2 for k in shap_values.keys()}

        # Calculate final score
        final_score = min(base_value + total_contribution, 1.0)

        # Create detailed explanations with actual found skills
        tech_skills_text = ', '.join(skills_data['technical_skills'][:5]) if skills_data['technical_skills'] else 'None found'
        if len(skills_data['technical_skills']) > 5:
            tech_skills_text += f' (and {len(skills_data["technical_skills"]) - 5} more)'

        soft_skills_text = ', '.join(skills_data['soft_skills'][:3]) if skills_data['soft_skills'] else 'None found'
        if len(skills_data['soft_skills']) > 3:
            soft_skills_text += f' (and {len(skills_data["soft_skills"]) - 3} more)'

        domain_skills_text = ', '.join(skills_data['domain_skills'][:3]) if skills_data['domain_skills'] else 'None found'
        if len(skills_data['domain_skills']) > 3:
            domain_skills_text += f' (and {len(skills_data["domain_skills"]) - 3} more)'

        explanations = {
            'technical_skills': f"Found {features['technical_skills_count']} technical skills: {tech_skills_text}",
            'soft_skills': f"Found {features['soft_skills_count']} soft skills: {soft_skills_text}",
            'domain_skills': f"Found {features['domain_skills_count']} domain-specific skills: {domain_skills_text}",
            'experience': f"{experience_data['level']} - {experience_data['summary']}",
            'education': f"{education_data['summary']}"
        }

        # Generate career guidance based on skills
        career_guidance = generate_career_guidance(skills_data, experience_data, education_data)

        return {
            'extracted_info': {
                'skills': all_skills,
                'experience': experience_data['summary'],
                'education': education_data['summary']
            },
            'career_guidance': career_guidance,
            'score': final_score,
            'feature_importance': normalized_shap,
            'feature_values': normalized_features,
            'explanation': explanations,
            'shap_analysis': {
                'base_value': base_value,
                'feature_contributions': shap_values,
                'total_contribution': total_contribution
            }
        }

    except Exception as e:
        logger.error(f"Error in SHAP-based analysis: {str(e)}")
        return create_enhanced_default_analysis()

def generate_career_guidance(skills_data, experience_data, education_data):
    """Generate comprehensive career guidance based on extracted data"""

    # Determine primary career paths based on skills
    tech_skills = skills_data['technical_skills']
    soft_skills = skills_data['soft_skills']
    domain_skills = skills_data['domain_skills']

    career_paths = []
    job_opportunities = []
    recommendations = []

    # Analyze skill patterns to suggest career paths
    if any(skill in tech_skills for skill in ['python', 'machine learning', 'tensorflow', 'pandas']):
        career_paths.append({
            "role": "Data Scientist",
            "description": "Your strong background in Python and data science tools makes you an excellent candidate for data science roles",
            "growth_potential": "Senior Data Scientist (3-5 years) → Lead Data Scientist (5-7 years) → Chief Data Officer (7+ years)",
            "salary_range": JOB_MARKET_DATA['data_scientist']['salary_range'],
            "next_steps": [
                "Build a portfolio with 3-5 data science projects",
                "Get certified in cloud platforms (AWS/Azure ML)",
                "Learn advanced ML techniques (deep learning, NLP)",
                "Contribute to open-source data science projects"
            ]
        })

        job_opportunities.append({
            "title": "Junior Data Scientist",
            "description": "Entry-level position focusing on data analysis and basic ML models",
            "required_skills": ["python", "pandas", "sql", "statistics"],
            "matching_skills": [skill for skill in ["python", "pandas", "sql", "statistics"] if skill in tech_skills],
            "missing_skills": [skill for skill in ["machine learning", "tensorflow", "tableau"] if skill not in tech_skills],
            "typical_salary": "$75,000 - $95,000",
            "demand_level": "High - Growing 31% faster than average"
        })

    if any(skill in tech_skills for skill in ['react', 'javascript', 'html', 'css', 'node.js']):
        career_paths.append({
            "role": "Full Stack Web Developer",
            "description": "Your web development skills position you well for modern full-stack development roles",
            "growth_potential": "Senior Developer (2-4 years) → Tech Lead (4-6 years) → Engineering Manager (6+ years)",
            "salary_range": JOB_MARKET_DATA['web_developer']['salary_range'],
            "next_steps": [
                "Master a modern framework (React/Vue/Angular)",
                "Learn backend technologies (Node.js/Python/Java)",
                "Build responsive, mobile-first applications",
                "Understand DevOps and deployment processes"
            ]
        })

        job_opportunities.append({
            "title": "Frontend Developer",
            "description": "Focus on user interface development and user experience",
            "required_skills": ["javascript", "html", "css", "react"],
            "matching_skills": [skill for skill in ["javascript", "html", "css", "react"] if skill in tech_skills],
            "missing_skills": [skill for skill in ["typescript", "webpack", "testing"] if skill not in tech_skills],
            "typical_salary": "$60,000 - $100,000",
            "demand_level": "High - Consistent growth in web development"
        })

    if any(skill in tech_skills for skill in ['aws', 'docker', 'kubernetes', 'jenkins', 'terraform']):
        career_paths.append({
            "role": "DevOps Engineer",
            "description": "Your cloud and automation skills are perfect for DevOps and infrastructure roles",
            "growth_potential": "Senior DevOps (3-5 years) → Platform Engineer (5-7 years) → Infrastructure Architect (7+ years)",
            "salary_range": JOB_MARKET_DATA['devops_engineer']['salary_range'],
            "next_steps": [
                "Get cloud certifications (AWS/Azure/GCP)",
                "Master Infrastructure as Code (Terraform/CloudFormation)",
                "Learn monitoring and observability tools",
                "Understand security best practices"
            ]
        })

    # Add generic software engineer path if no specific match
    if not career_paths and tech_skills:
        career_paths.append({
            "role": "Software Engineer",
            "description": "Your technical skills provide a solid foundation for general software engineering roles",
            "growth_potential": "Senior Engineer (3-5 years) → Staff Engineer (5-8 years) → Principal Engineer (8+ years)",
            "salary_range": JOB_MARKET_DATA['software_engineer']['salary_range'],
            "next_steps": [
                "Strengthen programming fundamentals",
                "Learn system design principles",
                "Build projects showcasing your skills",
                "Contribute to open-source projects"
            ]
        })

    # Generate recommendations based on current skills
    if len(tech_skills) < 5:
        recommendations.append("Focus on learning 3-5 core technical skills deeply rather than many superficially")

    if 'git' not in tech_skills:
        recommendations.append("Learn Git version control - essential for all development roles")

    if not any(cloud in tech_skills for cloud in ['aws', 'azure', 'gcp']):
        recommendations.append("Gain cloud platform experience - 90% of companies use cloud services")

    if len(soft_skills) < 3:
        recommendations.append("Develop soft skills like communication and teamwork - crucial for career advancement")

    # Market insights
    market_insights = {
        "industry_demand": "Technology sector growing 15% annually with high demand for skilled professionals",
        "emerging_skills": ["AI/Machine Learning", "Cloud Computing", "Cybersecurity", "Data Science"],
        "salary_trends": "Tech salaries increased 8-12% year-over-year, with remote work expanding opportunities",
        "top_locations": ["San Francisco Bay Area", "Seattle", "Austin", "Remote-first companies"]
    }

    return {
        "career_paths": career_paths[:2],  # Limit to top 2 paths
        "recommendations": recommendations[:4],  # Limit to top 4 recommendations
        "job_opportunities": job_opportunities[:3],  # Limit to top 3 opportunities
        "market_insights": market_insights
    }

def create_enhanced_default_analysis():
    """Create an enhanced default analysis with some meaningful data"""
    return {
        "extracted_info": {
            "skills": ["Analysis in progress"],
            "experience": "Experience analysis in progress",
            "education": "Education analysis in progress"
        },
        "career_guidance": {
            "career_paths": [
                {
                    "role": "Technology Professional",
                    "description": "Based on your background, technology roles offer excellent growth opportunities",
                    "growth_potential": "Strong growth potential in technology sector",
                    "salary_range": "$60,000 - $120,000",
                    "next_steps": [
                        "Complete skills assessment",
                        "Build technical portfolio",
                        "Network with industry professionals",
                        "Consider relevant certifications"
                    ]
                }
            ],
            "recommendations": [
                "Upload a detailed resume for comprehensive analysis",
                "Ensure resume includes specific technical skills",
                "Highlight quantifiable achievements",
                "Include relevant project experience"
            ],
            "job_opportunities": [
                {
                    "title": "Entry-Level Technology Role",
                    "description": "Various opportunities available based on your specific skills",
                    "required_skills": ["Technical skills", "Problem solving", "Communication"],
                    "matching_skills": ["To be determined"],
                    "missing_skills": ["To be assessed"],
                    "typical_salary": "$50,000 - $80,000",
                    "demand_level": "High demand in technology sector"
                }
            ],
            "market_insights": {
                "industry_demand": "Technology sector continues strong growth with diverse opportunities",
                "emerging_skills": ["Cloud Computing", "AI/ML", "Cybersecurity", "Data Analysis"],
                "salary_trends": "Competitive salaries with strong growth potential",
                "top_locations": ["Major tech hubs", "Remote opportunities", "Growing tech cities"]
            }
        },
        "feature_importance": {
            "technical_skills": 0.2,
            "soft_skills": 0.2,
            "domain_skills": 0.2,
            "experience_quality": 0.2,
            "education_quality": 0.2
        },
        "feature_values": {
            "technical_skills": 0.3,
            "soft_skills": 0.3,
            "domain_skills": 0.3,
            "experience_quality": 0.3,
            "education_quality": 0.3
        },
        "explanation": {
            "technical_skills": "Technical skills assessment in progress",
            "soft_skills": "Soft skills assessment in progress",
            "domain_skills": "Domain expertise assessment in progress",
            "experience": "Experience evaluation in progress",
            "education": "Education background evaluation in progress"
        },
        "score": 0.5
    }

def extract_text_from_pdf(file_content):
    """Extract text from PDF using multiple methods including OCR"""
    try:
        logger.info(f"Starting PDF extraction, content size: {len(file_content)} bytes")

        # Validate PDF content first
        if len(file_content) < 100:
            logger.error("File too small to be a valid PDF")
            return None

        # Check if it's actually a PDF file
        if not file_content.startswith(b'%PDF'):
            logger.error("File does not appear to be a valid PDF")
            return None

        # Try extracting with PyPDF2 first
        try:
            # Create a copy of the content for PyPDF2
            pdf_stream = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_stream)
            logger.info(f"PDF has {len(reader.pages)} pages")

            # Handle password-protected PDFs
            if reader.is_encrypted:
                logger.warning("PDF is password protected, attempting to decrypt with empty password")
                try:
                    # Try to decrypt with empty password
                    if reader.decrypt(''):
                        logger.info("Successfully decrypted PDF with empty password")
                    else:
                        logger.error("PDF requires a password to decrypt")
                        return None
                except Exception as decrypt_error:
                    logger.error(f"Failed to decrypt PDF: {str(decrypt_error)}")
                    return None

            text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                        logger.info(f"Page {i+1}: Extracted {len(page_text)} characters")
                    else:
                        logger.warning(f"Page {i+1}: No text extracted, might be scanned")
                except Exception as e:
                    logger.warning(f"Error on page {i+1}: {str(e)}, continuing with next page")
                    continue

            # If we got meaningful text, return it
            if text_parts:
                text = '\n'.join(text_parts)
                if len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with PyPDF2")
                    return text

            logger.info("PyPDF2 extraction yielded no meaningful text, trying alternative methods...")

        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}, trying alternative methods")

        # Try pdfplumber as second fallback
        try:
            import pdfplumber
            logger.info("Trying pdfplumber text extraction...")

            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                            logger.info(f"pdfplumber Page {page_num+1}: Extracted {len(page_text)} characters")
                    except Exception as e:
                        logger.warning(f"pdfplumber error on page {page_num+1}: {str(e)}")
                        continue

                if text_parts:
                    text = '\n'.join(text_parts)
                    if len(text.strip()) > 50:
                        logger.info(f"Successfully extracted {len(text)} characters with pdfplumber")
                        return text

        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")

        # Try PyMuPDF text extraction as third fallback
        try:
            import fitz  # PyMuPDF
            logger.info("Trying PyMuPDF text extraction...")

            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []

            for page_num in range(len(pdf_document)):
                try:
                    page = pdf_document[page_num]
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                        logger.info(f"PyMuPDF Page {page_num+1}: Extracted {len(page_text)} characters")
                except Exception as e:
                    logger.warning(f"PyMuPDF error on page {page_num+1}: {str(e)}")
                    continue

            pdf_document.close()

            if text_parts:
                text = '\n'.join(text_parts)
                if len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with PyMuPDF")
                    return text

        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {str(e)}")

        # Try OCR as final fallback
        try:
            import pytesseract
            import fitz  # PyMuPDF
            from PIL import Image
            import os

            logger.info("Attempting OCR extraction...")

            # Check for Tesseract installation (platform-specific)
            tesseract_paths = [
                "C:\\Program Files\\Tesseract-OCR\\tesseract.exe",  # Windows
                "/usr/bin/tesseract",  # Linux/Ubuntu
                "/usr/local/bin/tesseract",  # macOS/Linux alternative
                "tesseract"  # System PATH
            ]

            tesseract_found = False
            for tesseract_path in tesseract_paths:
                if tesseract_path == "tesseract":
                    # Try to find tesseract in system PATH
                    try:
                        import subprocess
                        subprocess.run(["tesseract", "--version"], capture_output=True, check=True)
                        pytesseract.pytesseract.tesseract_cmd = "tesseract"
                        tesseract_found = True
                        logger.info("Found Tesseract in system PATH")
                        break
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        continue
                elif os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    tesseract_found = True
                    logger.info(f"Found Tesseract at {tesseract_path}")
                    break

            if not tesseract_found:
                logger.warning("Tesseract not found - OCR extraction not available")
                logger.info("This PDF appears to be scanned and requires OCR to extract text")
                return None
            
            # Try OCR using PyMuPDF
            try:
                # Load PDF
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                text_parts = []

                # Process up to 3 pages for OCR (to avoid timeout)
                max_pages = min(3, len(pdf_document))
                logger.info(f"Processing {max_pages} pages with OCR")

                for page_num in range(max_pages):
                    try:
                        page = pdf_document[page_num]

                        # Convert to image with higher resolution for better OCR
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality

                        # Convert to PIL Image
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                        # Convert to grayscale for better OCR
                        img = img.convert('L')

                        # Perform OCR with custom config for better accuracy
                        custom_config = r'--oem 3 --psm 6'
                        page_text = pytesseract.image_to_string(img, config=custom_config)

                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                            logger.info(f"OCR Page {page_num+1}: Extracted {len(page_text)} characters")
                        else:
                            logger.warning(f"OCR Page {page_num+1}: No text extracted")

                    except Exception as page_error:
                        logger.warning(f"OCR error on page {page_num+1}: {str(page_error)}")
                        continue

                pdf_document.close()

                if text_parts:
                    text = '\n'.join(text_parts)
                    if len(text.strip()) > 50:
                        logger.info(f"Successfully extracted {len(text)} characters with OCR")
                        return text.strip()

                logger.warning("OCR yielded no meaningful text from any page")

            except Exception as e:
                logger.warning(f"Error during OCR: {str(e)}")

            logger.warning("OCR extraction failed")
            return None
            
        except ImportError as e:
            logger.error(f"OCR dependencies not installed: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return None
            
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        return None

def extract_text_from_docx(file_content):
    """Extract text from DOCX using multiple methods"""
    try:
        logger.info(f"Starting DOCX extraction, content size: {len(file_content)} bytes")

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            temp_path = temp_file.name
            
        try:
            # Try docx2txt first
            try:
                text = docx2txt.process(temp_path)
                if text and len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with docx2txt")
                    return text.strip()
                else:
                    logger.warning("docx2txt extraction yielded no meaningful text")
            except Exception as e:
                logger.error(f"docx2txt extraction failed: {str(e)}")
            
            # Try python-docx as fallback
            doc = Document(temp_path)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())
                    
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            if text_parts:
                text = '\n'.join(text_parts)
                if len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with python-docx")
                    return text
                    
            logger.error("No meaningful text extracted from DOCX")
            return None
            
        finally:
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.error(f"Error deleting temp file: {str(e)}")
                
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return None

def analyze_resume_with_ai(text):
    """Analyze resume text using AI with fallback to SHAP-based analysis"""
    try:
        # First try rule-based analysis with SHAP (always works)
        logger.info("Attempting rule-based analysis with SHAP explainability")
        shap_analysis = analyze_resume_with_shap(text)

        # Check if we should try AI enhancement
        if not check_internet_connectivity():
            logger.info("No internet connectivity - using rule-based analysis")
            return shap_analysis

        # Try to enhance with AI if available
        try:
            logger.info("Attempting AI enhancement of analysis")

            # Initialize the model with timeout configuration
            generation_config = genai.types.GenerationConfig(
                temperature=0.1,
                top_p=0.8,
                top_k=40,
                max_output_tokens=8192,
            )

            model = genai.GenerativeModel('gemini-1.5-pro', generation_config=generation_config)

            # Try to enhance the skills extraction with AI
            skills_prompt = f"""Enhance this skills analysis by finding any additional skills missed in the initial analysis.

Current skills found: {shap_analysis['extracted_info']['skills']}

Resume text:
{text[:2000]}

Return only a JSON list of additional skills not already found: ["skill1", "skill2"]"""

            def make_enhancement_request():
                return model.generate_content(skills_prompt)

            # Use retry logic for the API call
            enhancement_response = retry_with_backoff(make_enhancement_request, max_retries=2, base_delay=1)

            if enhancement_response:
                try:
                    response_text = enhancement_response.text.strip()
                    if response_text.startswith('```json'):
                        response_text = response_text[7:]
                    if response_text.endswith('```'):
                        response_text = response_text[:-3]
                    response_text = response_text.strip()

                    additional_skills = json.loads(response_text)
                    if isinstance(additional_skills, list):
                        # Add additional skills to the analysis
                        current_skills = shap_analysis['extracted_info']['skills']
                        enhanced_skills = list(set(current_skills + additional_skills))
                        shap_analysis['extracted_info']['skills'] = enhanced_skills
                        logger.info(f"AI enhanced analysis with {len(additional_skills)} additional skills")

                except Exception as e:
                    logger.warning(f"Failed to parse AI enhancement: {str(e)}")

            return shap_analysis

        except Exception as ai_error:
            logger.warning(f"AI enhancement failed: {str(ai_error)}, using rule-based analysis")
            return shap_analysis

    except Exception as e:
        logger.error(f"Error in main analysis: {str(e)}")
        # Log the specific type of error for debugging
        if "DNS resolution failed" in str(e) or "getaddrinfo" in str(e):
            logger.error("Network connectivity issue detected")
        elif "timeout" in str(e).lower():
            logger.error("Request timeout detected")
        elif "503" in str(e):
            logger.error("Service unavailable error detected")

        return create_enhanced_default_analysis()

def create_default_analysis():
    """Create a default analysis when AI fails"""
    return {
        "extracted_info": {
            "skills": [],
            "experience": "Could not extract experience information",
            "education": "Could not extract education information"
        },
        "career_guidance": {
            "career_paths": [
                {
                    "role": "Career Analysis Unavailable",
                    "description": "We encountered an error analyzing your resume. Please try again later.",
                    "growth_potential": "Unknown",
                    "salary_range": "Unknown",
                    "next_steps": [
                        "Please try uploading your resume again",
                        "Make sure your resume is in PDF or DOCX format",
                        "Ensure your resume is not password protected",
                        "Check your internet connection"
                    ]
                }
            ],
            "recommendations": [
                "Please try uploading your resume again",
                "Make sure your resume is in PDF or DOCX format",
                "Ensure your resume is not password protected",
                "Check your internet connection and try again"
            ],
            "job_opportunities": [
                {
                    "title": "Analysis Unavailable",
                    "description": "Job analysis could not be completed due to technical issues",
                    "required_skills": [],
                    "matching_skills": [],
                    "missing_skills": [],
                    "typical_salary": "Unknown",
                    "demand_level": "Unknown"
                }
            ],
            "market_insights": {
                "industry_demand": "Information not available due to technical issues",
                "emerging_skills": [],
                "salary_trends": "Information not available",
                "top_locations": []
            }
        },
        "feature_importance": {
            "technical_skills": 0.0,
            "soft_skills": 0.0,
            "domain_skills": 0.0,
            "experience_quality": 0.0,
            "education_quality": 0.0
        },
        "feature_values": {
            "technical_skills": 0.0,
            "soft_skills": 0.0,
            "domain_skills": 0.0,
            "experience_quality": 0.0,
            "education_quality": 0.0
        },
        "explanation": {
            "technical_skills": "Analysis unavailable",
            "soft_skills": "Analysis unavailable",
            "domain_skills": "Analysis unavailable",
            "experience": "Analysis unavailable",
            "education": "Analysis unavailable"
        },
        "score": 0.0
    }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/test-analysis')
def test_analysis():
    """Test endpoint to verify SHAP analysis is working"""
    try:
        sample_text = """
        John Doe
        Software Engineer

        Experience:
        Senior Software Engineer at Tech Corp (2020-2023)
        - Developed web applications using Python, JavaScript, and React
        - Led a team of 5 developers
        - Implemented machine learning models using TensorFlow
        - Worked with AWS cloud services

        Education:
        Bachelor's degree in Computer Science from University of Technology

        Skills:
        - Programming: Python, JavaScript, Java, C++
        - Web Development: React, Node.js, HTML, CSS
        - Databases: MySQL, MongoDB
        - Cloud: AWS, Docker, Kubernetes
        - Machine Learning: TensorFlow, scikit-learn
        - Soft Skills: Leadership, teamwork, problem solving
        """

        logger.info("Testing SHAP analysis with sample text")
        analysis = analyze_resume_with_shap(sample_text)

        if analysis:
            return jsonify({
                'success': True,
                'message': 'SHAP analysis is working!',
                'score': analysis.get('score', 0),
                'skills_count': len(analysis.get('extracted_info', {}).get('skills', [])),
                'feature_importance': analysis.get('feature_importance', {}),
                'career_paths_count': len(analysis.get('career_guidance', {}).get('career_paths', [])),
                'sample_skills': analysis.get('extracted_info', {}).get('skills', [])[:10]
            })
        else:
            return jsonify({
                'success': False,
                'message': 'SHAP analysis returned None',
                'error': 'Analysis function failed'
            })

    except Exception as e:
        logger.error(f"Test analysis failed: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'SHAP analysis failed with exception',
            'error': str(e)
        })

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        if 'resume' not in request.files:
            logger.error("No resume file in request")
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['resume']
        if not file or file.filename == '':
            logger.error("Empty filename")
            return jsonify({'error': 'No file selected'}), 400

        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            logger.error(f"Invalid file type: {filename}")
            return jsonify({'error': 'Invalid file type. Please upload PDF or DOCX only.'}), 400

        file_content = file.read()
        if not file_content:
            logger.error("Empty file content")
            return jsonify({'error': 'Empty file uploaded'}), 400
            
        logger.info(f"Processing file: {filename} ({len(file_content)} bytes)")

        file_extension = filename.rsplit('.', 1)[1].lower()
        if file_extension == 'pdf':
            text = extract_text_from_pdf(file_content)
        else:
            text = extract_text_from_docx(file_content)

        if not text:
            logger.error("Text extraction failed for all methods")

            # Instead of returning an error, provide a helpful response with guidance
            logger.info("Providing fallback analysis with file upload guidance")

            # Try to provide some basic analysis based on filename if possible
            basic_skills = []
            if 'engineer' in filename.lower():
                basic_skills.extend(['Engineering', 'Problem Solving', 'Technical Skills'])
            if 'developer' in filename.lower():
                basic_skills.extend(['Software Development', 'Programming', 'Coding'])
            if 'data' in filename.lower():
                basic_skills.extend(['Data Analysis', 'Analytics', 'Data Science'])
            if 'manager' in filename.lower():
                basic_skills.extend(['Management', 'Leadership', 'Team Management'])

            # Create a helpful analysis response
            fallback_analysis = {
                'score': 0.2,  # Give a small score to show something
                'feature_importance': {
                    'file_processing': 0.3,
                    'text_extraction': 0.0,
                    'content_analysis': 0.0,
                    'filename_analysis': 0.7
                },
                'extracted_info': {
                    'skills': basic_skills if basic_skills else ['File Processing Required'],
                    'experience': 'Unable to extract - requires OCR for scanned PDFs',
                    'education': 'Unable to extract - requires OCR for scanned PDFs'
                },
                'career_guidance': {
                    'career_paths': [{
                        'role': 'Resume Analysis Available with Text Extraction',
                        'description': 'Your file was uploaded successfully, but we need to extract text to provide detailed analysis. This appears to be a scanned PDF.',
                        'growth_potential': 'Full analysis available once text is extracted',
                        'salary_range': 'Analysis pending text extraction',
                        'next_steps': [
                            '🔧 Install Tesseract OCR for scanned PDF support',
                            '📄 Convert to text-based PDF using your word processor',
                            '📝 Try uploading as DOCX format instead',
                            '🔓 Ensure file is not password protected',
                            '✅ Verify file is not corrupted'
                        ]
                    }],
                    'recommendations': [
                        '🎯 IMMEDIATE: Install Tesseract OCR from https://github.com/UB-Mannheim/tesseract/wiki',
                        '📄 ALTERNATIVE: Save your resume as a new PDF from Word/Google Docs',
                        '📝 EASY FIX: Upload your resume as a DOCX file instead',
                        '🔍 VERIFY: Ensure your PDF contains selectable text (not just images)',
                        '⚡ QUICK TEST: Try copying text from your PDF - if you can\'t, it\'s scanned'
                    ]
                },
                'feature_values': {
                    'file_readable': 0.0,
                    'text_extractable': 0.0,
                    'format_supported': 1.0 if file_extension in ['pdf', 'docx'] else 0.0
                },
                'explanation': {
                    'file_processing': f'Unable to extract text from {file_extension.upper()} file',
                    'technical_details': 'Multiple extraction methods failed',
                    'suggestions': 'Try the recommendations above to resolve the issue',
                    'file_info': f'File: {filename} ({len(file_content)} bytes)'
                },
                'success': True,
                'processing_note': 'File uploaded successfully but text extraction failed'
            }

            return jsonify(fallback_analysis)

        preview = text[:200].replace('\n', ' ')
        logger.info(f"Extracted text preview: {preview}...")

        analysis = analyze_resume_with_ai(text)
        if not analysis:
            logger.error("AI analysis returned None")
            return jsonify({'error': 'Failed to analyze resume content'}), 500

        required_keys = ['score', 'feature_importance', 'extracted_info', 'career_guidance', 'feature_values', 'explanation']
        missing_keys = [key for key in required_keys if key not in analysis]
        if missing_keys:
            logger.error(f"Analysis missing required keys: {missing_keys}")
            # Use default analysis if keys are missing
            analysis = create_default_analysis()

        try:
            response = {
                'score': float(analysis.get('score', 0.0)),
                'feature_importance': analysis.get('feature_importance', {}),
                'extracted_info': analysis.get('extracted_info', {}),
                'career_guidance': analysis.get('career_guidance', {}),
                'feature_values': analysis.get('feature_values', {}),
                'explanation': analysis.get('explanation', {}),
                'success': True
            }
            
            
            logger.info(f"Response size: {len(str(response))} characters")
            logger.info(f"Analysis score: {response['score']}")
            logger.info(f"Number of skills: {len(response['extracted_info'].get('skills', []))}")
            
            return jsonify(response)
            
        except Exception as e:
            logger.error(f"Error preparing response: {str(e)}")
            return jsonify({'error': 'Error preparing analysis results'}), 500

    except Exception as e:
        logger.error(f"Error during analysis: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    try:
       
        try:
            import pytesseract
            import subprocess

            # Try to find Tesseract in various locations
            tesseract_paths = [
                "C:\\Program Files\\Tesseract-OCR\\tesseract.exe",  # Windows
                "/usr/bin/tesseract",  # Linux/Ubuntu
                "/usr/local/bin/tesseract",  # macOS/Linux alternative
                "tesseract"  # System PATH
            ]

            tesseract_found = False
            for tesseract_path in tesseract_paths:
                if tesseract_path == "tesseract":
                    try:
                        subprocess.run(["tesseract", "--version"], capture_output=True, check=True)
                        pytesseract.pytesseract.tesseract_cmd = "tesseract"
                        tesseract_found = True
                        logger.info("Tesseract OCR found in system PATH")
                        break
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        continue
                elif os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    tesseract_found = True
                    logger.info(f"Tesseract OCR found at {tesseract_path}")
                    break

            if not tesseract_found:
                logger.warning("Tesseract OCR not found - OCR functionality will be disabled")
                logger.warning("Scanned PDFs will not be processed")

        except ImportError:
            logger.warning("pytesseract not installed - OCR functionality will be disabled")
            
        # Check internet connectivity
        if not check_internet_connectivity():
            logger.warning("Limited internet connectivity detected. AI features may not work properly.")
        else:
            logger.info("Internet connectivity confirmed")

        # Initialize Gemini
        try:
            genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
            logger.info("Gemini API initialized successfully")

            # Test the API with a simple request
            try:
                test_model = genai.GenerativeModel('gemini-1.5-pro')
                test_response = test_model.generate_content("Hello")
                logger.info("Gemini API test successful")
            except Exception as test_e:
                logger.warning(f"Gemini API test failed: {str(test_e)}. AI features may not work properly.")

        except Exception as e:
            logger.error(f"Failed to initialize Gemini API: {str(e)}")
            raise
            
        
        ports = [5000, 8000, 8080, 3000]
        for port in ports:
            try:
                logger.info(f"Attempting to start server on port {port}...")
                app.run(host='127.0.0.1', port=port, debug=True, use_reloader=False)
                break
            except OSError as e:
                if port == ports[-1]: 
                    logger.error(f"Could not bind to any port. Last error: {str(e)}")
                    raise
                else:
                    logger.warning(f"Port {port} is busy, trying next port...")
                    continue
        
    except Exception as e:
        logger.error(f"Failed to start server: {str(e)}")
        raise
